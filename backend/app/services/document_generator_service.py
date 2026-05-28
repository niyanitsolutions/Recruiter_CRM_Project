"""HRM — Document Generator Service (PDF + DOCX via reportlab + python-docx)"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Optional

# ─── ReportLab ────────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, LETTER, LEGAL, A3, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
pt = 1  # reportlab uses points as native unit; 1 pt == 1 unit
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, Image as RLImage, Flowable,
)
from reportlab.platypus.flowables import KeepTogether
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.graphics.shapes import Drawing, String
from reportlab.lib.utils import ImageReader

# QR code (built into reportlab)
try:
    from reportlab.graphics.barcode import qr as rl_qr
    from reportlab.graphics import renderPDF as renderPDFModule
    _HAS_QR = True
except ImportError:
    _HAS_QR = False

# ─── python-docx ──────────────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# PIL for image handling (optional — gracefully degrades)
try:
    from PIL import Image as PILImage
    _HAS_PIL = True
except ImportError:
    _HAS_PIL = False


def _hex_to_rgb(hex_color: str) -> tuple:
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c * 2 for c in hex_color)
    try:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    except Exception:
        return (0, 0, 0)


def _hex_to_rl_color(hex_color: str):
    r, g, b = _hex_to_rgb(hex_color)
    return colors.Color(r / 255, g / 255, b / 255)


def _page_size_rl(size_str: str, orientation: str):
    base = {"A4": A4, "LETTER": LETTER, "LEGAL": LEGAL, "A3": A3}.get(size_str.upper(), A4)
    return landscape(base) if orientation.lower() == "landscape" else base


def _align_rl(align: str):
    return {"center": TA_CENTER, "right": TA_RIGHT, "justify": TA_JUSTIFY}.get(
        (align or "left").lower(), TA_LEFT
    )


def _strip_html(text: str) -> str:
    """Very simple HTML tag stripper for plain-text fallback."""
    return re.sub(r"<[^>]+>", "", text or "").strip()


# ─── Main Generator ───────────────────────────────────────────────────────────

class DocumentGeneratorService:
    """
    Generates PDF or DOCX from a resolved document template.
    Input: the template dict (from DocumentTemplateService.generate) plus field_data.
    """

    # ─── Public entry points ──────────────────────────────────────────────────

    def generate_pdf(self, template: dict, resolved_blocks: list, field_data: dict) -> bytes:
        """Return a PDF as bytes."""
        buf = io.BytesIO()
        self._build_pdf(buf, template, resolved_blocks, field_data)
        buf.seek(0)
        return buf.read()

    def generate_docx(self, template: dict, resolved_blocks: list, field_data: dict) -> bytes:
        """Return a DOCX as bytes."""
        if not _HAS_DOCX:
            raise RuntimeError("python-docx is not installed.")
        buf = io.BytesIO()
        self._build_docx(buf, template, resolved_blocks, field_data)
        buf.seek(0)
        return buf.read()

    # ─── PDF Builder ──────────────────────────────────────────────────────────

    def _build_pdf(self, buf: io.BytesIO, template: dict, blocks: list, fields: dict) -> None:
        branding   = template.get("branding", {})
        page_cfg   = template.get("page_config", {})
        header_cfg = template.get("header", {})
        footer_cfg = template.get("footer", {})
        watermark  = template.get("watermark", {})

        page_size  = _page_size_rl(page_cfg.get("size", "A4"), page_cfg.get("orientation", "portrait"))
        mt = page_cfg.get("margin_top",    20) * mm
        mr = page_cfg.get("margin_right",  20) * mm
        mb = page_cfg.get("margin_bottom", 20) * mm
        ml = page_cfg.get("margin_left",   20) * mm

        primary_hex = branding.get("primary_color", "#1e3a5f")
        primary_clr = _hex_to_rl_color(primary_hex)
        text_hex    = branding.get("text_color", "#1a1a1a")
        heading_hex = branding.get("heading_color", "#1e3a5f")
        font_size   = branding.get("font_size", 11)
        font_family = "Helvetica"   # safe default; no custom fonts without TTF registration

        # Base styles
        sample_styles = getSampleStyleSheet()
        normal_style  = ParagraphStyle(
            "DocNormal",
            parent=sample_styles["Normal"],
            fontSize=font_size,
            textColor=_hex_to_rl_color(text_hex),
            leading=font_size * 1.5,
            spaceAfter=6,
        )
        h1_style = ParagraphStyle(
            "DocH1", parent=sample_styles["Heading1"],
            fontSize=font_size + 8, textColor=_hex_to_rl_color(heading_hex),
            fontName="Helvetica-Bold", spaceAfter=8,
        )
        h2_style = ParagraphStyle(
            "DocH2", parent=sample_styles["Heading2"],
            fontSize=font_size + 4, textColor=_hex_to_rl_color(heading_hex),
            fontName="Helvetica-Bold", spaceAfter=6,
        )
        h3_style = ParagraphStyle(
            "DocH3", parent=sample_styles["Heading3"],
            fontSize=font_size + 2, textColor=_hex_to_rl_color(heading_hex),
            fontName="Helvetica-Bold", spaceAfter=4,
        )

        # Callbacks for header / footer / watermark
        def _on_page(canvas_obj, doc):
            canvas_obj.saveState()

            if watermark.get("enabled"):
                self._draw_watermark_pdf(canvas_obj, page_size, watermark)

            if header_cfg.get("enabled", True):
                self._draw_header_pdf(canvas_obj, page_size, header_cfg, branding, fields, ml, mr, mt, page_cfg)

            if footer_cfg.get("enabled", True):
                self._draw_footer_pdf(canvas_obj, page_size, footer_cfg, primary_clr, ml, mr, mb)

            canvas_obj.restoreState()

        doc = SimpleDocTemplate(
            buf,
            pagesize=page_size,
            topMargin=mt + (30 * mm if header_cfg.get("enabled", True) else 0),
            bottomMargin=mb + (15 * mm if footer_cfg.get("enabled", True) else 0),
            leftMargin=ml,
            rightMargin=mr,
        )

        story = self._blocks_to_flowables(blocks, fields, branding, normal_style, h1_style, h2_style, h3_style, primary_clr)

        # Signature flowables
        sigs = template.get("signatures", [])
        if sigs:
            story.extend(self._signatures_to_flowables(sigs, normal_style, primary_clr))

        doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)

    def _draw_watermark_pdf(self, canvas_obj, page_size, watermark: dict) -> None:
        w, h = page_size
        canvas_obj.saveState()
        canvas_obj.setFillAlpha(watermark.get("opacity", 0.10))
        wm_text = watermark.get("text", "CONFIDENTIAL")
        wm_size = watermark.get("font_size", 60)
        rotation = watermark.get("rotation", -45)
        color_hex = watermark.get("color", "#cccccc")
        r, g, b = _hex_to_rgb(color_hex)
        canvas_obj.setFillColorRGB(r / 255, g / 255, b / 255)
        canvas_obj.setFont("Helvetica-Bold", wm_size)
        canvas_obj.translate(w / 2, h / 2)
        canvas_obj.rotate(rotation)
        canvas_obj.drawCentredString(0, 0, wm_text)
        canvas_obj.restoreState()

    def _draw_header_pdf(self, canvas_obj, page_size, header: dict, branding: dict,
                          fields: dict, ml, mr, mt, page_cfg) -> None:
        w, h = page_size
        header_y  = h - mt + 5 * mm
        header_h  = 28 * mm
        x_start   = ml
        x_end     = w - mr
        avail_w   = x_end - x_start

        # Background
        bg_hex = header.get("background_color", "#ffffff")
        r, g, b = _hex_to_rgb(bg_hex)
        canvas_obj.setFillColorRGB(r / 255, g / 255, b / 255)
        canvas_obj.rect(x_start, header_y - header_h, avail_w, header_h, fill=1, stroke=0)

        # Border bottom
        if header.get("border_bottom", True):
            primary_hex = branding.get("primary_color", "#1e3a5f")
            r2, g2, b2 = _hex_to_rgb(primary_hex)
            canvas_obj.setStrokeColorRGB(r2 / 255, g2 / 255, b2 / 255)
            canvas_obj.setLineWidth(1.5)
            canvas_obj.line(x_start, header_y - header_h, x_end, header_y - header_h)

        # Company name
        primary_hex = branding.get("primary_color", "#1e3a5f")
        r3, g3, b3 = _hex_to_rgb(primary_hex)
        canvas_obj.setFillColorRGB(r3 / 255, g3 / 255, b3 / 255)
        canvas_obj.setFont("Helvetica-Bold", 14)
        co_name = fields.get("company_name", "")
        if header.get("show_company_name") and co_name:
            canvas_obj.drawString(x_start, header_y - 12 * mm, co_name)

        # Company address / details
        canvas_obj.setFillColorRGB(0.3, 0.3, 0.3)
        canvas_obj.setFont("Helvetica", 8)
        detail_lines = []
        if header.get("show_address") and fields.get("company_address"):
            detail_lines.append(fields["company_address"])
        if header.get("show_phone") and fields.get("company_phone"):
            detail_lines.append(f"Tel: {fields['company_phone']}")
        if header.get("show_email") and fields.get("company_email"):
            detail_lines.append(f"Email: {fields['company_email']}")
        if header.get("show_website") and fields.get("company_website"):
            detail_lines.append(fields["company_website"])

        for i, line in enumerate(detail_lines[:3]):
            canvas_obj.drawString(x_start, header_y - (15 + i * 5) * mm, line)

    def _draw_footer_pdf(self, canvas_obj, page_size, footer: dict,
                          primary_clr, ml, mr, mb) -> None:
        w, h = page_size
        footer_y = mb - 2 * mm
        x_start  = ml
        x_end    = w - mr

        if footer.get("border_top", True):
            canvas_obj.setStrokeColor(primary_clr)
            canvas_obj.setLineWidth(0.5)
            canvas_obj.line(x_start, footer_y + 8 * mm, x_end, footer_y + 8 * mm)

        canvas_obj.setFillColorRGB(0.5, 0.5, 0.5)
        canvas_obj.setFont("Helvetica", 8)

        parts = []
        if footer.get("show_generated_date"):
            parts.append(f"Generated: {datetime.now().strftime('%d %b %Y')}")
        if footer.get("disclaimer"):
            parts.append(footer["disclaimer"][:80])

        footer_text = "  |  ".join(parts)
        if footer_text:
            canvas_obj.drawCentredString(w / 2, footer_y, footer_text)

        if footer.get("show_page_numbers"):
            canvas_obj.drawRightString(x_end, footer_y, "Page 1")

    # ─── Blocks → Flowables ───────────────────────────────────────────────────

    def _blocks_to_flowables(
        self, blocks: list, fields: dict, branding: dict,
        normal_style, h1, h2, h3, primary_clr
    ) -> list:
        story = []
        text_clr = _hex_to_rl_color(branding.get("text_color", "#1a1a1a"))
        primary_hex = branding.get("primary_color", "#1e3a5f")

        for block in sorted(blocks, key=lambda b: b.get("order", 0)):
            btype   = block.get("type", "text")
            content = block.get("content", "")
            props   = block.get("properties", {})

            mt = props.get("margin_top", 4)
            mb = props.get("margin_bottom", 4)
            fs = props.get("font_size") or branding.get("font_size", 11)
            align = _align_rl(props.get("text_align", "left"))
            color_hex = props.get("color") or branding.get("text_color", "#1a1a1a")
            color = _hex_to_rl_color(color_hex)

            if mt:
                story.append(Spacer(1, mt))

            if btype == "heading":
                tag = h1 if fs >= 20 else h2 if fs >= 16 else h3
                style = ParagraphStyle(
                    f"h_{id(block)}", parent=tag,
                    fontSize=fs, alignment=align, textColor=color,
                )
                safe_content = _strip_html(str(content)) if content else ""
                story.append(Paragraph(safe_content, style))

            elif btype in ("text", "paragraph"):
                style = ParagraphStyle(
                    f"p_{id(block)}", parent=normal_style,
                    fontSize=fs, alignment=align, textColor=color,
                )
                # Convert simple HTML to reportlab markup
                safe_content = self._html_to_rl(str(content))
                story.append(Paragraph(safe_content, style))

            elif btype == "divider":
                color_h = props.get("color", "#e2e8f0")
                story.append(HRFlowable(
                    width="100%", thickness=1,
                    color=_hex_to_rl_color(color_h),
                    spaceAfter=4,
                ))

            elif btype == "spacer":
                h_val = int(str(props.get("height", "20")).replace("px", "") or 20)
                story.append(Spacer(1, h_val))

            elif btype == "page_break":
                story.append(PageBreak())

            elif btype == "list_items":
                items = content if isinstance(content, list) else [str(content)]
                for item in items:
                    bullet_style = ParagraphStyle(
                        f"li_{id(item)}", parent=normal_style,
                        leftIndent=16, bulletIndent=4, fontSize=fs,
                    )
                    story.append(Paragraph(f"• {_strip_html(str(item))}", bullet_style))

            elif btype == "table":
                fl = self._table_to_flowable(content, props, primary_hex, normal_style, fs)
                if fl:
                    story.append(fl)

            elif btype == "salary_table":
                fl = self._salary_table_to_flowable(content, primary_hex, normal_style, fs)
                if fl:
                    story.append(fl)

            elif btype in ("employee_details", "company_details"):
                fl = self._kv_table_to_flowable(content, primary_hex, normal_style, fs)
                if fl:
                    story.append(fl)

            elif btype == "qr_code":
                if _HAS_QR and content:
                    try:
                        qr_obj  = rl_qr.QrCodeWidget(str(content))
                        bounds  = qr_obj.getBounds()
                        qr_w = bounds[2] - bounds[0]
                        qr_h = bounds[3] - bounds[1]
                        d = Drawing(60, 60, transform=[60./qr_w, 0, 0, 60./qr_h, 0, 0])
                        d.add(qr_obj)
                        story.append(d)
                    except Exception:
                        story.append(Paragraph(f"[QR: {content}]", normal_style))

            elif btype == "two_column":
                if isinstance(content, dict):
                    left_text  = _strip_html(str(content.get("left", "")))
                    right_text = _strip_html(str(content.get("right", "")))
                    tbl = Table(
                        [[Paragraph(left_text, normal_style), Paragraph(right_text, normal_style)]],
                        colWidths=["50%", "50%"],
                    )
                    tbl.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
                    story.append(tbl)

            if mb:
                story.append(Spacer(1, mb))

        return story

    def _html_to_rl(self, html: str) -> str:
        """Convert basic HTML markup to reportlab Paragraph XML."""
        if not html:
            return ""
        # reportlab Paragraph supports <b>, <i>, <u>, <br/>, &amp; etc.
        # Strip unsafe tags; keep formatting
        html = re.sub(r"<(div|p|span)[^>]*>", "", html)
        html = re.sub(r"</(div|p|span)>", "<br/>", html)
        html = re.sub(r"<br\s*/?>", "<br/>", html)
        # Keep <b>, <i>, <u>, <strong>, <em>
        html = re.sub(r"<strong>", "<b>", html)
        html = re.sub(r"</strong>", "</b>", html)
        html = re.sub(r"<em>", "<i>", html)
        html = re.sub(r"</em>", "</i>", html)
        # Strip remaining tags (reportlab will reject unknown ones)
        html = re.sub(r"<(?!/?[biuBIU]|br/)([^>]+)>", "", html)
        return html.strip()

    def _table_to_flowable(self, content, props: dict, primary_hex: str, normal_style, fs: int):
        if not isinstance(content, dict):
            return None
        headers  = content.get("headers", [])
        rows     = content.get("rows", [])
        has_hdr  = content.get("has_header", True)
        hdr_bg   = content.get("header_bg", primary_hex)
        hdr_clr  = content.get("header_color", "#ffffff")
        stripe   = content.get("stripe_rows", True)

        cell_style = ParagraphStyle("cell", parent=normal_style, fontSize=fs - 1, leading=fs * 1.3)

        table_data = []
        if has_hdr and headers:
            hdr_row = [Paragraph(str(h), ParagraphStyle("hdr", parent=cell_style, textColor=colors.white, fontName="Helvetica-Bold")) for h in headers]
            table_data.append(hdr_row)

        for row in rows:
            table_data.append([Paragraph(str(c), cell_style) for c in row])

        if not table_data:
            return None

        tbl = Table(table_data, repeatRows=1 if (has_hdr and headers) else 0)
        r1, g1, b1 = _hex_to_rgb(hdr_bg)
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0 if (has_hdr and headers) else -1),
             colors.Color(r1 / 255, g1 / 255, b1 / 255)),
            ("TEXTCOLOR",  (0, 0), (-1, 0),   colors.white),
            ("FONTNAME",   (0, 0), (-1, 0),   "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1),  fs - 1),
            ("GRID",       (0, 0), (-1, -1),  0.5, colors.Color(0.85, 0.85, 0.85)),
            ("VALIGN",     (0, 0), (-1, -1),  "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1),  4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 8),
        ]
        if stripe and has_hdr and headers:
            for i in range(1, len(table_data)):
                if i % 2 == 0:
                    style_cmds.append(("BACKGROUND", (0, i), (-1, i), colors.Color(0.97, 0.97, 0.97)))
        tbl.setStyle(TableStyle(style_cmds))
        return tbl

    def _salary_table_to_flowable(self, content, primary_hex: str, normal_style, fs: int):
        if not isinstance(content, dict):
            return None

        earnings   = content.get("earnings", [])
        deductions = content.get("deductions", [])

        def _sum(items):
            total = 0.0
            for item in items:
                try:
                    total += float(str(item.get("value", 0)).replace(",", "").replace("₹", "").strip() or 0)
                except (ValueError, TypeError):
                    pass
            return total

        gross  = _sum(earnings)
        deduct = _sum(deductions)
        net    = gross - deduct

        r1, g1, b1 = _hex_to_rgb(primary_hex)
        primary_rl  = colors.Color(r1 / 255, g1 / 255, b1 / 255)
        light_bg    = colors.Color(0.97, 0.98, 1.0)

        cell_s = ParagraphStyle("sc", parent=normal_style, fontSize=fs - 1, leading=fs * 1.3)
        bold_s = ParagraphStyle("scb", parent=cell_s, fontName="Helvetica-Bold")
        white_s = ParagraphStyle("scw", parent=bold_s, textColor=colors.white)

        earn_rows   = [[Paragraph(e["label"], cell_s), Paragraph(f"₹{e.get('value','')}", cell_s)] for e in earnings]
        deduct_rows = [[Paragraph(d["label"], cell_s), Paragraph(f"₹{d.get('value','')}", cell_s)] for d in deductions]

        max_rows = max(len(earn_rows), len(deduct_rows))
        while len(earn_rows)   < max_rows: earn_rows.append([Paragraph("", cell_s), Paragraph("", cell_s)])
        while len(deduct_rows) < max_rows: deduct_rows.append([Paragraph("", cell_s), Paragraph("", cell_s)])

        data = [
            [Paragraph("Earnings", white_s), Paragraph("Amount", white_s),
             Paragraph("Deductions", white_s), Paragraph("Amount", white_s)],
        ]
        for er, dr in zip(earn_rows, deduct_rows):
            data.append([er[0], er[1], dr[0], dr[1]])

        data.append([
            Paragraph("Gross Salary", bold_s), Paragraph(f"₹{gross:,.2f}", bold_s),
            Paragraph("Total Deductions", bold_s), Paragraph(f"₹{deduct:,.2f}", bold_s),
        ])
        data.append([
            Paragraph("NET PAY", ParagraphStyle("np", parent=bold_s, textColor=colors.white)),
            Paragraph(f"₹{net:,.2f}", ParagraphStyle("npa", parent=bold_s, textColor=colors.white)),
            Paragraph("", white_s),
            Paragraph("", white_s),
        ])

        tbl = Table(data, colWidths=["25%", "25%", "25%", "25%"])
        n = len(data)
        tbl.setStyle(TableStyle([
            ("BACKGROUND",  (0, 0), (-1, 0),     primary_rl),
            ("TEXTCOLOR",   (0, 0), (-1, 0),     colors.white),
            ("FONTNAME",    (0, 0), (-1, 0),     "Helvetica-Bold"),
            ("BACKGROUND",  (0, n - 2), (-1, n - 2), light_bg),
            ("FONTNAME",    (0, n - 2), (-1, n - 2), "Helvetica-Bold"),
            ("BACKGROUND",  (0, n - 1), (-1, n - 1), primary_rl),
            ("SPAN",        (0, n - 1), (1, n - 1)),
            ("GRID",        (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.85, 0.85)),
            ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",  (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ]))
        return tbl

    def _kv_table_to_flowable(self, content, primary_hex: str, normal_style, fs: int):
        if not isinstance(content, dict):
            return None
        r1, g1, b1 = _hex_to_rgb(primary_hex)
        primary_rl = colors.Color(r1 / 255, g1 / 255, b1 / 255)
        label_s = ParagraphStyle("kl", parent=normal_style, fontSize=fs - 1, fontName="Helvetica-Bold", textColor=primary_rl)
        val_s   = ParagraphStyle("kv", parent=normal_style, fontSize=fs - 1)
        data = [[Paragraph(k.replace("_", " ").title(), label_s), Paragraph(str(v), val_s)]
                for k, v in content.items() if v]
        if not data:
            return None
        tbl = Table(data, colWidths=["35%", "65%"])
        tbl.setStyle(TableStyle([
            ("GRID",        (0, 0), (-1, -1), 0.5, colors.Color(0.88, 0.88, 0.88)),
            ("VALIGN",      (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",  (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ]))
        return tbl

    def _signatures_to_flowables(self, signatures: list, normal_style, primary_clr) -> list:
        if not signatures:
            return []
        story = [Spacer(1, 20), HRFlowable(width="100%", thickness=0.5, color=primary_clr)]
        sig_cells = []
        for sig in signatures:
            label = sig.get("label", "Authorized Signatory")
            name  = sig.get("name", "")
            desig = sig.get("designation", "")
            sig_block = [
                Spacer(1, 35),
                HRFlowable(width=100, thickness=1, color=colors.black),
                Paragraph(f"<b>{name or label}</b>", normal_style),
            ]
            if desig:
                sig_block.append(Paragraph(desig, ParagraphStyle("sd", parent=normal_style, fontSize=9, textColor=colors.gray)))
            sig_block.append(Paragraph(label, ParagraphStyle("sl", parent=normal_style, fontSize=8, textColor=colors.lightgrey)))
            sig_cells.append(sig_block)

        if sig_cells:
            tbl = Table([sig_cells], colWidths=[f"{100 // len(sig_cells)}%" for _ in sig_cells])
            tbl.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "BOTTOM"), ("ALIGN", (0, 0), (-1, -1), "LEFT")]))
            story.append(tbl)
        return story

    # ─── DOCX Builder ─────────────────────────────────────────────────────────

    def _build_docx(self, buf: io.BytesIO, template: dict, blocks: list, fields: dict) -> None:
        doc = DocxDocument()
        branding   = template.get("branding", {})
        header_cfg = template.get("header", {})
        footer_cfg = template.get("footer", {})

        primary_hex = branding.get("primary_color", "#1e3a5f")
        font_family = branding.get("font_family", "Calibri")
        font_size   = branding.get("font_size", 11)
        text_hex    = branding.get("text_color", "#1a1a1a")
        heading_hex = branding.get("heading_color", "#1e3a5f")

        # Page margins
        page_cfg = template.get("page_config", {})
        for section in doc.sections:
            section.top_margin    = Cm(page_cfg.get("margin_top", 20) / 10)
            section.bottom_margin = Cm(page_cfg.get("margin_bottom", 20) / 10)
            section.left_margin   = Cm(page_cfg.get("margin_left", 20) / 10)
            section.right_margin  = Cm(page_cfg.get("margin_right", 20) / 10)

        # Header
        if header_cfg.get("enabled", True):
            hdr = doc.sections[0].header
            hp  = hdr.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if header_cfg.get("show_company_name") and fields.get("company_name"):
                run = hp.add_run(fields["company_name"])
                run.bold  = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(*_hex_to_rgb(primary_hex))
            if header_cfg.get("show_address") and fields.get("company_address"):
                hp.add_run(f"\n{fields['company_address']}")

        # Footer
        if footer_cfg.get("enabled", True):
            ftr = doc.sections[0].footer
            fp  = ftr.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parts = []
            if footer_cfg.get("show_generated_date"):
                parts.append(f"Generated: {datetime.now().strftime('%d %b %Y')}")
            if footer_cfg.get("disclaimer"):
                parts.append(footer_cfg["disclaimer"][:80])
            if parts:
                fp.add_run("  |  ".join(parts)).font.size = Pt(8)

        # Content blocks
        for block in sorted(blocks, key=lambda b: b.get("order", 0)):
            self._docx_add_block(doc, block, branding, fields)

        # Signatures
        sigs = template.get("signatures", [])
        if sigs:
            doc.add_paragraph()
            doc.add_paragraph("─" * 60)
            for sig in sigs:
                p = doc.add_paragraph()
                r = p.add_run(f"\n\n{'_' * 30}")
                p.add_run(f"\n{sig.get('name', sig.get('label', 'Signatory'))}")
                p.add_run(f"\n{sig.get('designation', '')}")
                p.add_run(f"\n{sig.get('label', 'Authorized Signatory')}")

        doc.save(buf)

    def _docx_add_block(self, doc: "DocxDocument", block: dict, branding: dict, fields: dict) -> None:
        btype   = block.get("type", "text")
        content = block.get("content", "")
        props   = block.get("properties", {})

        primary_hex = branding.get("primary_color", "#1e3a5f")
        font_family = branding.get("font_family", "Calibri")
        fs          = props.get("font_size") or branding.get("font_size", 11)
        align_str   = (props.get("text_align") or "left").lower()
        align_map   = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                       "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
        align = align_map.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)

        if btype == "heading":
            level = 1 if fs >= 20 else 2 if fs >= 16 else 3
            p = doc.add_heading(_strip_html(str(content)), level=level)
            p.alignment = align
            for run in p.runs:
                run.font.color.rgb = RGBColor(*_hex_to_rgb(branding.get("heading_color", primary_hex)))

        elif btype in ("text", "paragraph"):
            p = doc.add_paragraph(_strip_html(str(content)))
            p.alignment = align
            for run in p.runs:
                run.font.name = font_family
                run.font.size = Pt(fs)

        elif btype == "divider":
            p = doc.add_paragraph("─" * 80)
            for run in p.runs:
                run.font.size = Pt(6)
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        elif btype == "spacer":
            doc.add_paragraph()

        elif btype == "page_break":
            doc.add_page_break()

        elif btype == "list_items":
            items = content if isinstance(content, list) else [str(content)]
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(_strip_html(str(item))).font.size = Pt(fs)

        elif btype == "table":
            if isinstance(content, dict):
                headers = content.get("headers", [])
                rows    = content.get("rows", [])
                has_hdr = content.get("has_header", True)
                total_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
                num_rows = (1 if (has_hdr and headers) else 0) + len(rows)
                if num_rows > 0 and total_cols > 0:
                    tbl = doc.add_table(rows=num_rows, cols=total_cols)
                    tbl.style = "Table Grid"
                    row_idx = 0
                    if has_hdr and headers:
                        for ci, h in enumerate(headers):
                            cell = tbl.rows[row_idx].cells[ci]
                            cell.text = str(h)
                            r, g, b = _hex_to_rgb(primary_hex)
                            self._docx_set_cell_bg(cell, primary_hex.lstrip("#"))
                            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(h))
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        row_idx += 1
                    for row in rows:
                        for ci, cell_val in enumerate(row):
                            if ci < total_cols:
                                tbl.rows[row_idx].cells[ci].text = str(cell_val)
                        row_idx += 1

        elif btype == "salary_table":
            if isinstance(content, dict):
                earnings   = content.get("earnings", [])
                deductions = content.get("deductions", [])
                all_rows   = [["Earnings", "Amount", "Deductions", "Amount"]]
                max_r = max(len(earnings), len(deductions))
                for i in range(max_r):
                    e = earnings[i]   if i < len(earnings)   else {"label": "", "value": ""}
                    d = deductions[i] if i < len(deductions) else {"label": "", "value": ""}
                    all_rows.append([e.get("label", ""), f"₹{e.get('value','')}", d.get("label", ""), f"₹{d.get('value','')}"])
                if all_rows:
                    tbl = doc.add_table(rows=len(all_rows), cols=4)
                    tbl.style = "Table Grid"
                    for ri, row in enumerate(all_rows):
                        for ci, val in enumerate(row):
                            cell = tbl.rows[ri].cells[ci]
                            cell.text = str(val)
                            if ri == 0:
                                self._docx_set_cell_bg(cell, primary_hex.lstrip("#"))
                                for r in cell.paragraphs[0].runs:
                                    r.font.color.rgb = RGBColor(255, 255, 255)
                                    r.bold = True

        elif btype in ("employee_details", "company_details"):
            if isinstance(content, dict):
                pairs = [(k.replace("_", " ").title(), str(v)) for k, v in content.items() if v]
                if pairs:
                    tbl = doc.add_table(rows=len(pairs), cols=2)
                    tbl.style = "Table Grid"
                    for ri, (k, v) in enumerate(pairs):
                        tbl.rows[ri].cells[0].text = k
                        tbl.rows[ri].cells[1].text = v
                        for run in tbl.rows[ri].cells[0].paragraphs[0].runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(*_hex_to_rgb(primary_hex))

        elif btype == "two_column":
            if isinstance(content, dict):
                left  = _strip_html(str(content.get("left", "")))
                right = _strip_html(str(content.get("right", "")))
                tbl = doc.add_table(rows=1, cols=2)
                tbl.rows[0].cells[0].text = left
                tbl.rows[0].cells[1].text = right

    @staticmethod
    def _docx_set_cell_bg(cell, hex_color: str) -> None:
        """Set table cell background color via direct XML manipulation."""
        try:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_color.upper())
            tcPr.append(shd)
        except Exception:
            pass
