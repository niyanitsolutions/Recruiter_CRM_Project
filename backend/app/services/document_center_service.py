"""
Document Center Service
Handles templates, versioning, PDF/DOCX generation, approvals.
"""
from datetime import datetime, timezone
from typing import Optional, List, Tuple, Dict, Any
import uuid
import io
import re
import logging
from html.parser import HTMLParser

from motor.motor_asyncio import AsyncIOMotorDatabase
from bson import ObjectId

from app.models.company.document_center import (
    DocCategory, DocTemplate, DocTemplateVersion, DocGenerated, DocApproval,
    DocCategoryCreate, DocCategoryUpdate,
    DocTemplateCreate, DocTemplateUpdate,
    DocGenerateRequest, DocSendRequest,
    DocApprovalCreate, DocApprovalReview,
    DocTemplateType, DocStatus, ApprovalStatus,
    TemplateContent, PaperSize, PaperOrientation,
)

logger = logging.getLogger(__name__)

# ─── Pre-built template library ───────────────────────────────────────────────

TEMPLATE_LIBRARY: List[Dict[str, Any]] = [
    {"key": "offer_letter",       "name": "Offer Letter",        "category": "HR Letters",   "description": "Standard employment offer letter"},
    {"key": "appointment_letter", "name": "Appointment Letter",  "category": "HR Letters",   "description": "Formal appointment confirmation"},
    {"key": "joining_letter",     "name": "Joining Letter",      "category": "HR Letters",   "description": "Welcome letter for new joiners"},
    {"key": "experience_letter",  "name": "Experience Letter",   "category": "HR Letters",   "description": "Work experience certificate"},
    {"key": "relieving_letter",   "name": "Relieving Letter",    "category": "HR Letters",   "description": "Employee relieving confirmation"},
    {"key": "promotion_letter",   "name": "Promotion Letter",    "category": "HR Letters",   "description": "Promotion announcement letter"},
    {"key": "transfer_letter",    "name": "Transfer Letter",     "category": "HR Letters",   "description": "Employee transfer order"},
    {"key": "warning_letter",     "name": "Warning Letter",      "category": "HR Letters",   "description": "Official warning notice"},
    {"key": "termination_letter", "name": "Termination Letter",  "category": "HR Letters",   "description": "Employment termination notice"},
    {"key": "nda",                "name": "NDA Agreement",       "category": "Legal",        "description": "Non-disclosure agreement"},
    {"key": "employment_agreement","name": "Employment Agreement","category": "Legal",        "description": "Full employment contract"},
    {"key": "payslip",            "name": "Payslip",             "category": "Payroll",      "description": "Monthly salary slip"},
    {"key": "salary_certificate", "name": "Salary Certificate",  "category": "Payroll",      "description": "Salary verification certificate"},
    {"key": "internship_letter",  "name": "Internship Letter",   "category": "HR Letters",   "description": "Internship confirmation letter"},
    {"key": "asset_handover",     "name": "Asset Handover",      "category": "Operations",   "description": "Asset handover acknowledgement"},
    {"key": "exit_clearance",     "name": "Exit Clearance",      "category": "Operations",   "description": "Exit clearance checklist"},
]

TEMPLATE_BODIES: Dict[str, str] = {
    "offer_letter": """<h2>Offer Letter</h2>
<p>Date: {{current_date}}</p>
<p>To,<br>{{employee_name}}<br>{{employee_address}}</p>
<p>Dear {{employee_name}},</p>
<p>We are pleased to offer you the position of <strong>{{designation}}</strong> in the <strong>{{department}}</strong> department at <strong>{{company_name}}</strong>.</p>
<h3>Terms & Conditions</h3>
<ul>
<li><strong>Designation:</strong> {{designation}}</li>
<li><strong>Department:</strong> {{department}}</li>
<li><strong>Date of Joining:</strong> {{joining_date}}</li>
<li><strong>CTC:</strong> {{salary}} per annum</li>
<li><strong>Reporting To:</strong> {{manager_name}}</li>
</ul>
<p>Please confirm your acceptance of this offer by signing and returning a copy of this letter.</p>
<p>We look forward to welcoming you to the team.</p>
<p>Yours sincerely,<br><strong>HR Department</strong><br>{{company_name}}</p>""",

    "experience_letter": """<h2>Experience Certificate</h2>
<p>Date: {{current_date}}</p>
<p>To Whom It May Concern,</p>
<p>This is to certify that <strong>{{employee_name}}</strong> (Employee ID: {{employee_id}}) was employed with <strong>{{company_name}}</strong> as <strong>{{designation}}</strong> in the <strong>{{department}}</strong> department from <strong>{{joining_date}}</strong> to <strong>{{exit_date}}</strong>.</p>
<p>During their tenure, {{employee_name}} demonstrated dedication, professionalism, and a strong work ethic. We wish them the very best in their future endeavours.</p>
<p>Yours faithfully,<br><strong>HR Department</strong><br>{{company_name}}</p>""",

    "relieving_letter": """<h2>Relieving Letter</h2>
<p>Date: {{current_date}}</p>
<p>To,<br>{{employee_name}}</p>
<p>Dear {{employee_name}},</p>
<p>This is to confirm that you have been relieved of your duties as <strong>{{designation}}</strong> at <strong>{{company_name}}</strong> effective <strong>{{exit_date}}</strong>.</p>
<p>All company assets, documents, and credentials in your possession have been returned in satisfactory condition. We confirm that there are no pending dues on either side.</p>
<p>We thank you for your valuable contribution and wish you the very best in your future career.</p>
<p>Yours sincerely,<br><strong>HR Department</strong><br>{{company_name}}</p>""",

    "payslip": """<h2>Salary Slip — {{month_year}}</h2>
<table>
<tr><th colspan="2">Employee Information</th><th colspan="2">Pay Period</th></tr>
<tr><td>Name</td><td>{{employee_name}}</td><td>Month</td><td>{{month_year}}</td></tr>
<tr><td>Employee ID</td><td>{{employee_id}}</td><td>Paid Days</td><td>{{paid_days}}</td></tr>
<tr><td>Designation</td><td>{{designation}}</td><td>LOP Days</td><td>{{lop_days}}</td></tr>
<tr><td>Department</td><td>{{department}}</td><td>Bank</td><td>{{bank_name}}</td></tr>
</table>
<br>
<table>
<tr><th>Earnings</th><th>Amount</th><th>Deductions</th><th>Amount</th></tr>
<tr><td>Basic Salary</td><td>{{basic}}</td><td>Provident Fund</td><td>{{pf}}</td></tr>
<tr><td>HRA</td><td>{{hra}}</td><td>Professional Tax</td><td>{{pt}}</td></tr>
<tr><td>Special Allowance</td><td>{{special_allowance}}</td><td>TDS</td><td>{{tds}}</td></tr>
<tr><td><strong>Gross Salary</strong></td><td><strong>{{gross}}</strong></td><td><strong>Total Deductions</strong></td><td><strong>{{total_deductions}}</strong></td></tr>
</table>
<p><strong>Net Pay: {{net_salary}}</strong></p>
<p><em>Amount in Words: {{salary_in_words}}</em></p>""",

    "warning_letter": """<h2>Warning Letter</h2>
<p>Date: {{current_date}}</p>
<p>To,<br>{{employee_name}}<br>{{designation}}, {{department}}</p>
<p>Dear {{employee_name}},</p>
<p>This letter serves as a formal warning regarding your conduct/performance at <strong>{{company_name}}</strong>.</p>
<p><strong>Nature of Issue:</strong><br>{{issue_description}}</p>
<p><strong>Date of Incident:</strong> {{incident_date}}</p>
<p>This behavior is in violation of company policy and is not acceptable. We expect immediate improvement and compliance with all company policies and procedures.</p>
<p>Failure to improve may result in further disciplinary action, including termination of employment.</p>
<p>Please acknowledge receipt of this letter by signing below.</p>
<p>Yours sincerely,<br><strong>HR Manager</strong><br>{{company_name}}</p>""",

    "nda": """<h2>Non-Disclosure Agreement</h2>
<p>This Agreement is entered into as of <strong>{{current_date}}</strong> between:</p>
<p><strong>{{company_name}}</strong> (the "Company"), and<br><strong>{{employee_name}}</strong> (the "Employee")</p>
<h3>1. Confidential Information</h3>
<p>The Employee agrees to hold in strict confidence all proprietary and confidential information of the Company, including but not limited to business strategies, client lists, technical data, and trade secrets.</p>
<h3>2. Non-Disclosure</h3>
<p>The Employee shall not disclose any Confidential Information to any third party without prior written consent from the Company.</p>
<h3>3. Term</h3>
<p>This Agreement shall remain in effect during the Employee's employment and for a period of <strong>2 years</strong> after termination.</p>
<h3>4. Return of Information</h3>
<p>Upon termination, the Employee shall return all materials containing Confidential Information to the Company.</p>
<p>IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.</p>
<p>___________________________&nbsp;&nbsp;&nbsp;___________________________<br>
{{employee_name}}&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;For {{company_name}}<br>
Date:___________&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;Date:___________</p>""",
}

# ─── HTML Text Extractor ───────────────────────────────────────────────────────

class _HtmlTextExtractor(HTMLParser):
    """Extracts structured text blocks from HTML for PDF rendering."""

    def __init__(self):
        super().__init__()
        self.blocks: List[Dict[str, Any]] = []
        self._current_tag: str = ""
        self._current_text: List[str] = []
        self._in_table: bool = False
        self._table_rows: List[List[str]] = []
        self._current_row: List[str] = []
        self._current_cell: List[str] = []
        self._list_items: List[str] = []
        self._in_list: bool = False
        self._list_type: str = "ul"

    def handle_starttag(self, tag, attrs):
        self._current_tag = tag
        if tag in ("table",):
            self._in_table = True
            self._table_rows = []
        elif tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._current_cell = []
        elif tag in ("ul", "ol"):
            self._in_list = True
            self._list_type = tag
            self._list_items = []
        elif tag == "li":
            self._current_cell = []

    def handle_endtag(self, tag):
        text = " ".join(self._current_text).strip()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            if text:
                self.blocks.append({"type": "heading", "level": int(tag[1]), "text": text})
            self._current_text = []
        elif tag == "p":
            if text:
                self.blocks.append({"type": "paragraph", "text": text})
            self._current_text = []
        elif tag in ("td", "th"):
            cell_text = " ".join(self._current_cell).strip()
            self._current_row.append(cell_text)
            self._current_cell = []
        elif tag == "tr":
            if self._current_row:
                self._table_rows.append(self._current_row)
                self._current_row = []
        elif tag == "table":
            if self._table_rows:
                self.blocks.append({"type": "table", "rows": self._table_rows})
            self._in_table = False
            self._table_rows = []
        elif tag == "li":
            item_text = " ".join(self._current_cell).strip()
            if item_text:
                self._list_items.append(item_text)
            self._current_cell = []
        elif tag in ("ul", "ol"):
            if self._list_items:
                self.blocks.append({"type": "list", "list_type": self._list_type, "items": self._list_items})
            self._in_list = False
            self._list_items = []
        self._current_tag = ""

    def handle_data(self, data):
        clean = data.strip()
        if not clean:
            return
        if self._in_table and self._current_tag in ("td", "th"):
            self._current_cell.append(clean)
        elif self._in_list and self._current_tag == "li":
            self._current_cell.append(clean)
        else:
            self._current_text.append(clean)


def _extract_blocks(html: str) -> List[Dict[str, Any]]:
    parser = _HtmlTextExtractor()
    try:
        parser.feed(html)
    except Exception:
        pass
    return parser.blocks or [{"type": "paragraph", "text": re.sub(r"<[^>]+>", " ", html).strip()}]


def _strip_html(html: str) -> str:
    return re.sub(r"<[^>]+>", " ", html).strip()


# ─── PDF Generator ────────────────────────────────────────────────────────────

def _generate_pdf(template: DocTemplate, html_content: str, field_values: Dict[str, str]) -> bytes:
    """
    Generate a PDF from a template using ReportLab.
    Returns PDF bytes.
    """
    from reportlab.lib.pagesizes import A4, letter, legal, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import pt, inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table as RLTable,
        TableStyle, HRFlowable, PageBreak,
    )
    from reportlab.platypus.flowables import Flowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

    content_obj = template.content
    paper_cfg   = content_obj.paper
    header_cfg  = content_obj.header
    footer_cfg  = content_obj.footer
    wm_cfg      = content_obj.watermark

    # Page size
    size_map = {
        PaperSize.A4:     A4,
        PaperSize.LETTER: letter,
        PaperSize.LEGAL:  legal,
    }
    page_size = size_map.get(paper_cfg.size, A4)
    if paper_cfg.orientation == PaperOrientation.LANDSCAPE:
        page_size = landscape(page_size)

    buf = io.BytesIO()

    margin_t = paper_cfg.margin_top   * pt
    margin_b = paper_cfg.margin_bottom * pt
    margin_l = paper_cfg.margin_left  * pt
    margin_r = paper_cfg.margin_right * pt

    # Build header/footer callbacks
    def _on_page(canvas, doc):
        canvas.saveState()

        # Watermark
        if wm_cfg.enabled and wm_cfg.type.value == "text" and wm_cfg.text:
            canvas.setFont("Helvetica-Bold", wm_cfg.size)
            canvas.setFillColorRGB(0.5, 0.5, 0.5, alpha=wm_cfg.opacity)
            cx = page_size[0] / 2
            cy = page_size[1] / 2
            canvas.translate(cx, cy)
            canvas.rotate(wm_cfg.rotation)
            canvas.drawCentredString(0, 0, wm_cfg.text)
            canvas.translate(-cx, -cy)

        # Header
        if header_cfg.show:
            hh_pt = getattr(header_cfg, 'header_height', 120) * 0.75  # px to pt approx
            header_y = page_size[1] - margin_t + hh_pt * 0.5 + 4
            canvas.setFont("Helvetica-Bold", header_cfg.font_size + 2)
            try:
                r, g, b = int(header_cfg.font_color[1:3], 16)/255, int(header_cfg.font_color[3:5], 16)/255, int(header_cfg.font_color[5:7], 16)/255
                canvas.setFillColorRGB(r, g, b)
            except Exception:
                canvas.setFillColorRGB(0, 0, 0)
            # Use company_alignment (new field) with fallback to alignment (old field)
            h_align = getattr(header_cfg, 'company_alignment', None) or getattr(header_cfg, 'alignment', 'left')
            align_x = {
                "left":   margin_l,
                "center": page_size[0] / 2,
                "right":  page_size[0] - margin_r,
            }.get(h_align, margin_l)
            draw_fn = {
                "left":   canvas.drawString,
                "center": canvas.drawCentredString,
                "right":  canvas.drawRightString,
            }.get(h_align, canvas.drawString)
            if header_cfg.company_name:
                draw_fn(align_x, header_y, header_cfg.company_name)
            if header_cfg.company_address:
                canvas.setFont("Helvetica", max(6, header_cfg.font_size - 1))
                draw_fn(align_x, header_y - 14, header_cfg.company_address)
            if header_cfg.border_bottom:
                border_y = page_size[1] - margin_t
                canvas.setStrokeColorRGB(0.7, 0.7, 0.7)
                canvas.line(margin_l, border_y, page_size[0] - margin_r, border_y)

        # Footer
        if footer_cfg.show:
            footer_y = margin_b - 20
            canvas.setFont("Helvetica", footer_cfg.font_size)
            try:
                r, g, b = int(footer_cfg.font_color[1:3], 16)/255, int(footer_cfg.font_color[3:5], 16)/255, int(footer_cfg.font_color[5:7], 16)/255
                canvas.setFillColorRGB(r, g, b)
            except Exception:
                canvas.setFillColorRGB(0.4, 0.4, 0.4)
            if footer_cfg.border_top:
                canvas.setStrokeColorRGB(0.7, 0.7, 0.7)
                canvas.line(margin_l, footer_y + 10, page_size[0] - margin_r, footer_y + 10)
            center_x = page_size[0] / 2
            footer_parts = []
            if footer_cfg.text:
                footer_parts.append(footer_cfg.text)
            if footer_cfg.confidential_label:
                footer_parts.append("CONFIDENTIAL")
            if footer_parts:
                canvas.drawCentredString(center_x, footer_y, "  |  ".join(footer_parts))
            if footer_cfg.show_page_numbers:
                page_text = f"Page {doc.page}"
                canvas.drawRightString(page_size[0] - margin_r, footer_y, page_text)
            if footer_cfg.show_date:
                canvas.drawString(margin_l, footer_y, datetime.now().strftime("%B %d, %Y"))

        canvas.restoreState()

    doc = SimpleDocTemplate(
        buf,
        pagesize=page_size,
        topMargin=margin_t,
        bottomMargin=margin_b,
        leftMargin=margin_l,
        rightMargin=margin_r,
    )

    styles   = getSampleStyleSheet()
    story    = []
    align_map = {"left": TA_LEFT, "center": TA_CENTER, "right": TA_RIGHT, "justify": TA_JUSTIFY}

    heading_sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 10}
    blocks = _extract_blocks(html_content)

    for block in blocks:
        if block["type"] == "heading":
            lvl = block.get("level", 1)
            sz = heading_sizes.get(lvl, 14)
            style = ParagraphStyle(
                name=f"H{lvl}",
                parent=styles["Normal"],
                fontSize=sz,
                fontName="Helvetica-Bold",
                spaceAfter=8,
                spaceBefore=12,
            )
            story.append(Paragraph(block["text"], style))

        elif block["type"] == "paragraph":
            style = ParagraphStyle(
                name="Body",
                parent=styles["Normal"],
                fontSize=11,
                leading=16,
                spaceAfter=6,
            )
            story.append(Paragraph(block["text"], style))

        elif block["type"] == "list":
            for i, item in enumerate(block["items"]):
                bullet = "•" if block["list_type"] == "ul" else f"{i+1}."
                style = ParagraphStyle(
                    name="ListItem",
                    parent=styles["Normal"],
                    fontSize=11,
                    leading=15,
                    leftIndent=20,
                    spaceAfter=3,
                )
                story.append(Paragraph(f"{bullet}  {item}", style))
            story.append(Spacer(1, 6))

        elif block["type"] == "table":
            rows_data = block["rows"]
            if not rows_data:
                continue
            col_count = max(len(r) for r in rows_data)
            # Normalize rows
            table_data = [
                [Paragraph(cell, styles["Normal"]) for cell in (row + [""] * col_count)[:col_count]]
                for row in rows_data
            ]
            col_width = (page_size[0] - margin_l - margin_r) / col_count
            tbl = RLTable(table_data, colWidths=[col_width] * col_count)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#7c3aed")),
                ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",   (0, 0), (-1, -1), 10),
                ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
                ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING",   (0, 0), (-1, -1), 8),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 10))

        story.append(Spacer(1, 4))

    if not story:
        story.append(Paragraph("No content.", styles["Normal"]))

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    return buf.getvalue()


# ─── DOCX Generator ───────────────────────────────────────────────────────────

def _generate_docx(template: DocTemplate, html_content: str, field_values: Dict[str, str]) -> bytes:
    """Generate a DOCX from a template using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    content_obj = template.content
    header_cfg  = content_obj.header
    footer_cfg  = content_obj.footer
    paper_cfg   = content_obj.paper

    doc = Document()

    # Page margins
    margin_in = paper_cfg.margin_top / 72.0
    for section in doc.sections:
        section.top_margin    = Inches(margin_in)
        section.bottom_margin = Inches(paper_cfg.margin_bottom / 72.0)
        section.left_margin   = Inches(paper_cfg.margin_left / 72.0)
        section.right_margin  = Inches(paper_cfg.margin_right / 72.0)

    # Header
    if header_cfg.show and header_cfg.company_name:
        for section in doc.sections:
            hdr = section.header
            p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            p.text = header_cfg.company_name
            if header_cfg.company_address:
                p.add_run(f"\n{header_cfg.company_address}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header_cfg.alignment == "center" else WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0] if p.runs else None
            if run:
                run.bold = True
                run.font.size = Pt(header_cfg.font_size + 2)

    # Footer
    if footer_cfg.show:
        for section in doc.sections:
            ftr = section.footer
            p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            footer_text = footer_cfg.text or ""
            if footer_cfg.confidential_label:
                footer_text = (footer_text + "  |  CONFIDENTIAL").strip(" |")
            p.text = footer_text
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Body content
    blocks = _extract_blocks(html_content)
    heading_styles = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}

    for block in blocks:
        if block["type"] == "heading":
            lvl  = block.get("level", 1)
            style = heading_styles.get(lvl, "Heading 3")
            p = doc.add_paragraph(block["text"], style=style)

        elif block["type"] == "paragraph":
            p = doc.add_paragraph(block["text"])
            p.style = doc.styles["Normal"]

        elif block["type"] == "list":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Bullet" if block["list_type"] == "ul" else "List Number")
                p.add_run(item)

        elif block["type"] == "table":
            rows_data = block["rows"]
            if not rows_data:
                continue
            col_count = max(len(r) for r in rows_data)
            tbl = doc.add_table(rows=len(rows_data), cols=col_count)
            tbl.style = "Table Grid"
            for ri, row in enumerate(rows_data):
                for ci, cell_text in enumerate(row[:col_count]):
                    cell = tbl.cell(ri, ci)
                    cell.text = cell_text
                    if ri == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Field Resolution ─────────────────────────────────────────────────────────

def _resolve_fields(html: str, values: Dict[str, str]) -> str:
    """Replace {{field}} placeholders with resolved values."""
    def replacer(match):
        key = match.group(1).strip()
        return values.get(key, match.group(0))
    return re.sub(r"\{\{(\w+)\}\}", replacer, html)


async def _fetch_employee_fields(db: AsyncIOMotorDatabase, employee_id: str) -> Dict[str, str]:
    """Load employee data from DB and map to template field names."""
    fields: Dict[str, str] = {}
    if not employee_id:
        return fields
    emp = await db.employees.find_one({"_id": employee_id, "is_deleted": False})
    if not emp:
        return fields
    now = datetime.now(timezone.utc)
    fields.update({
        "employee_name":    emp.get("full_name", ""),
        "employee_id":      emp.get("employee_id", emp.get("_id", "")),
        "department":       emp.get("department", ""),
        "designation":      emp.get("designation", ""),
        "salary":           str(emp.get("salary", "")),
        "joining_date":     str(emp.get("joining_date", "") or ""),
        "exit_date":        str(emp.get("exit_date", "") or ""),
        "manager_name":     emp.get("manager_name", ""),
        "employee_email":   emp.get("email", ""),
        "employee_address": emp.get("address", ""),
        "employee_phone":   emp.get("phone", ""),
        "current_date":     now.strftime("%B %d, %Y"),
        "month_year":       now.strftime("%B %Y"),
    })
    return fields


# ─── Service Class ─────────────────────────────────────────────────────────────

class DocumentCenterService:

    # ── Category CRUD ──────────────────────────────────────────────────────────

    async def list_categories(self, db: AsyncIOMotorDatabase) -> List[Dict]:
        cursor = db.doc_categories.find({"is_deleted": False}).sort("sort_order", 1)
        cats = await cursor.to_list(length=500)
        # Enrich with template counts
        for cat in cats:
            cat["template_count"] = await db.doc_templates.count_documents({
                "category_id": cat["_id"],
                "is_deleted": False,
            })
        return cats

    async def create_category(self, db: AsyncIOMotorDatabase, data: DocCategoryCreate, user_id: str) -> Dict:
        cat = DocCategory(
            name=data.name,
            description=data.description,
            color=data.color,
            icon=data.icon,
            sort_order=data.sort_order,
            created_by=user_id,
        )
        doc = cat.model_dump(by_alias=True)
        await db.doc_categories.insert_one(doc)
        return doc

    async def update_category(self, db: AsyncIOMotorDatabase, category_id: str, data: DocCategoryUpdate) -> Tuple[bool, str]:
        updates = {k: v for k, v in data.model_dump(exclude_none=True).items()}
        if not updates:
            return False, "No changes provided"
        updates["updated_at"] = datetime.now(timezone.utc)
        result = await db.doc_categories.update_one({"_id": category_id}, {"$set": updates})
        if result.matched_count == 0:
            return False, "Category not found"
        return True, "Category updated"

    async def delete_category(self, db: AsyncIOMotorDatabase, category_id: str) -> Tuple[bool, str]:
        # Unlink templates
        await db.doc_templates.update_many({"category_id": category_id}, {"$set": {"category_id": None, "category_name": ""}})
        result = await db.doc_categories.update_one({"_id": category_id}, {"$set": {"is_deleted": True}})
        if result.matched_count == 0:
            return False, "Category not found"
        return True, "Category deleted"

    # ── Template CRUD ──────────────────────────────────────────────────────────

    async def list_templates(
        self,
        db: AsyncIOMotorDatabase,
        *,
        category_id: Optional[str] = None,
        status: Optional[str] = None,
        template_type: Optional[str] = None,
        is_favorite: Optional[bool] = None,
        is_archived: Optional[bool] = None,
        search: Optional[str] = None,
        tags: Optional[List[str]] = None,
        skip: int = 0,
        limit: int = 50,
    ) -> Tuple[List[Dict], int]:
        query: Dict[str, Any] = {"is_deleted": False}
        if category_id:
            query["category_id"] = category_id
        if status:
            query["status"] = status
        if template_type:
            query["template_type"] = template_type
        if is_favorite is not None:
            query["is_favorite"] = is_favorite
        if is_archived is not None:
            query["is_archived"] = is_archived
        else:
            query["is_archived"] = {"$ne": True}
        if search:
            query["$or"] = [
                {"name":        {"$regex": search, "$options": "i"}},
                {"description": {"$regex": search, "$options": "i"}},
                {"tags":        {"$elemMatch": {"$regex": search, "$options": "i"}}},
            ]
        if tags:
            query["tags"] = {"$all": tags}

        total = await db.doc_templates.count_documents(query)
        # Project out heavy content field for list view
        projection = {"content.body_html": 0, "content.canvas_elements": 0, "content.custom_css": 0}
        cursor = db.doc_templates.find(query, projection).sort("updated_at", -1).skip(skip).limit(limit)
        docs = await cursor.to_list(length=limit)

        # Attach category color
        cat_ids = list({d.get("category_id") for d in docs if d.get("category_id")})
        cat_map: Dict[str, Dict] = {}
        if cat_ids:
            cats = await db.doc_categories.find({"_id": {"$in": cat_ids}}).to_list(length=500)
            cat_map = {c["_id"]: c for c in cats}
        for d in docs:
            cid = d.get("category_id")
            d["category_color"] = cat_map[cid]["color"] if cid and cid in cat_map else None

        return docs, total

    async def get_template(self, db: AsyncIOMotorDatabase, template_id: str) -> Optional[Dict]:
        doc = await db.doc_templates.find_one({"_id": template_id, "is_deleted": False})
        if doc:
            cat = await db.doc_categories.find_one({"_id": doc.get("category_id")}) if doc.get("category_id") else None
            doc["category_color"] = cat["color"] if cat else None
        return doc

    async def create_template(
        self, db: AsyncIOMotorDatabase, data: DocTemplateCreate, user_id: str, user_name: str
    ) -> Dict:
        # Resolve category name
        cat_name = ""
        if data.category_id:
            cat = await db.doc_categories.find_one({"_id": data.category_id})
            cat_name = cat["name"] if cat else ""

        content = data.content or TemplateContent()
        tmpl = DocTemplate(
            name=data.name,
            description=data.description,
            category_id=data.category_id,
            category_name=cat_name,
            template_type=data.template_type,
            tags=data.tags,
            content=content,
            dynamic_fields=data.dynamic_fields,
            created_by=user_id,
            created_by_name=user_name,
            updated_by=user_id,
        )
        doc = tmpl.model_dump(by_alias=True)
        await db.doc_templates.insert_one(doc)

        # Save initial version
        await self._save_version(db, doc["_id"], 1, data.name, content, data.change_summary, user_id, user_name)
        return doc

    async def update_template(
        self, db: AsyncIOMotorDatabase, template_id: str, data: DocTemplateUpdate, user_id: str, user_name: str
    ) -> Tuple[bool, str]:
        existing = await db.doc_templates.find_one({"_id": template_id, "is_deleted": False})
        if not existing:
            return False, "Template not found"

        updates: Dict[str, Any] = {}
        if data.name         is not None: updates["name"]         = data.name
        if data.description  is not None: updates["description"]  = data.description
        if data.category_id  is not None:
            updates["category_id"] = data.category_id
            cat = await db.doc_categories.find_one({"_id": data.category_id})
            updates["category_name"] = cat["name"] if cat else ""
        if data.status       is not None: updates["status"]       = data.status
        if data.tags         is not None: updates["tags"]         = data.tags
        if data.is_favorite  is not None: updates["is_favorite"]  = data.is_favorite
        if data.is_archived  is not None: updates["is_archived"]  = data.is_archived
        if data.dynamic_fields is not None: updates["dynamic_fields"] = data.dynamic_fields
        if data.content      is not None:
            updates["content"]    = data.content.model_dump()
            new_version = existing.get("version", 1) + 1
            updates["version"]    = new_version
            await self._save_version(
                db, template_id, new_version,
                data.name or existing["name"],
                data.content,
                data.change_summary,
                user_id, user_name,
            )

        if not updates:
            return True, "No changes"
        updates["updated_at"] = datetime.now(timezone.utc)
        updates["updated_by"] = user_id
        await db.doc_templates.update_one({"_id": template_id}, {"$set": updates})
        return True, "Template updated"

    async def duplicate_template(
        self, db: AsyncIOMotorDatabase, template_id: str, user_id: str, user_name: str
    ) -> Tuple[bool, str, Optional[Dict]]:
        existing = await db.doc_templates.find_one({"_id": template_id, "is_deleted": False})
        if not existing:
            return False, "Template not found", None
        content_obj = TemplateContent(**existing.get("content", {}))
        tmpl = DocTemplate(
            name=f"Copy of {existing['name']}",
            description=existing.get("description", ""),
            category_id=existing.get("category_id"),
            category_name=existing.get("category_name", ""),
            template_type=DocTemplateType(existing.get("template_type", "simple")),
            tags=existing.get("tags", []),
            content=content_obj,
            dynamic_fields=existing.get("dynamic_fields", []),
            created_by=user_id,
            created_by_name=user_name,
            updated_by=user_id,
        )
        doc = tmpl.model_dump(by_alias=True)
        await db.doc_templates.insert_one(doc)
        await self._save_version(
            db, doc["_id"], 1, doc["name"], content_obj,
            f"Duplicated from {existing['name']}", user_id, user_name,
        )
        return True, "Template duplicated", doc

    async def delete_template(self, db: AsyncIOMotorDatabase, template_id: str) -> Tuple[bool, str]:
        result = await db.doc_templates.update_one(
            {"_id": template_id},
            {"$set": {"is_deleted": True, "updated_at": datetime.now(timezone.utc)}},
        )
        if result.matched_count == 0:
            return False, "Template not found"
        return True, "Template deleted"

    async def toggle_favorite(self, db: AsyncIOMotorDatabase, template_id: str) -> Tuple[bool, str, bool]:
        doc = await db.doc_templates.find_one({"_id": template_id})
        if not doc:
            return False, "Template not found", False
        new_val = not doc.get("is_favorite", False)
        await db.doc_templates.update_one({"_id": template_id}, {"$set": {"is_favorite": new_val}})
        return True, "Favorite updated", new_val

    # ── Version History ────────────────────────────────────────────────────────

    async def _save_version(
        self, db, template_id: str, version: int, name: str,
        content: TemplateContent, summary: str, user_id: str, user_name: str,
    ) -> None:
        ver = DocTemplateVersion(
            template_id=template_id,
            version=version,
            name=name,
            content=content,
            change_summary=summary,
            created_by=user_id,
            created_by_name=user_name,
        )
        await db.doc_template_versions.insert_one(ver.model_dump(by_alias=True))

    async def list_versions(self, db: AsyncIOMotorDatabase, template_id: str) -> List[Dict]:
        cursor = db.doc_template_versions.find({"template_id": template_id}).sort("version", -1)
        return await cursor.to_list(length=200)

    async def restore_version(
        self, db: AsyncIOMotorDatabase, template_id: str, version_id: str, user_id: str, user_name: str
    ) -> Tuple[bool, str]:
        ver = await db.doc_template_versions.find_one({"_id": version_id, "template_id": template_id})
        if not ver:
            return False, "Version not found"
        tmpl = await db.doc_templates.find_one({"_id": template_id})
        if not tmpl:
            return False, "Template not found"
        new_version = tmpl.get("version", 1) + 1
        content_obj = TemplateContent(**ver["content"])
        updates = {
            "content": ver["content"],
            "version": new_version,
            "updated_at": datetime.now(timezone.utc),
            "updated_by": user_id,
        }
        await db.doc_templates.update_one({"_id": template_id}, {"$set": updates})
        await self._save_version(
            db, template_id, new_version, tmpl["name"],
            content_obj, f"Restored from version {ver['version']}",
            user_id, user_name,
        )
        return True, f"Restored to version {ver['version']}"

    async def delete_version(self, db: AsyncIOMotorDatabase, template_id: str, version_id: str) -> Tuple[bool, str]:
        # Can't delete latest version
        tmpl = await db.doc_templates.find_one({"_id": template_id})
        if not tmpl:
            return False, "Template not found"
        ver = await db.doc_template_versions.find_one({"_id": version_id, "template_id": template_id})
        if not ver:
            return False, "Version not found"
        if ver.get("version") == tmpl.get("version"):
            return False, "Cannot delete the current version"
        await db.doc_template_versions.delete_one({"_id": version_id})
        return True, "Version deleted"

    # ── Document Generation ────────────────────────────────────────────────────

    async def generate_document(
        self,
        db: AsyncIOMotorDatabase,
        req: DocGenerateRequest,
        user_id: str,
        user_name: str,
    ) -> Tuple[bool, str, Optional[Dict]]:
        tmpl = await db.doc_templates.find_one({"_id": req.template_id, "is_deleted": False})
        if not tmpl:
            return False, "Template not found", None

        # Build field values: employee DB → user-supplied overrides
        resolved: Dict[str, str] = {"current_date": datetime.now().strftime("%B %d, %Y")}
        if req.employee_id:
            emp_fields = await _fetch_employee_fields(db, req.employee_id)
            resolved.update(emp_fields)

        # Fetch company name from tenant settings if available
        settings = await db.company_settings.find_one({}) or {}
        company_name = settings.get("company_name", "")
        if company_name:
            resolved["company_name"] = company_name

        # User overrides last
        resolved.update(req.field_values)

        # Build rendered HTML
        body_html = tmpl.get("content", {}).get("body_html", "")
        html_content = _resolve_fields(body_html, resolved)

        # Reconstruct template object for PDF/DOCX generation
        content_data = tmpl.get("content", {})
        content_obj = TemplateContent(**content_data)
        tmpl_obj = DocTemplate(**{**tmpl, "content": content_obj})

        pdf_bytes  = None
        docx_bytes = None
        pdf_url    = None
        docx_url   = None

        if req.generate_pdf:
            try:
                pdf_bytes = _generate_pdf(tmpl_obj, html_content, resolved)
            except Exception as e:
                logger.error(f"PDF generation failed: {e}")

        if req.generate_docx:
            try:
                docx_bytes = _generate_docx(tmpl_obj, html_content, resolved)
            except Exception as e:
                logger.error(f"DOCX generation failed: {e}")

        # Store to S3 if S3 is configured
        try:
            from app.utils.s3 import upload_bytes
            doc_id = str(uuid.uuid4())
            if pdf_bytes:
                key = f"doc-center/generated/{doc_id}.pdf"
                pdf_url = await upload_bytes(key, pdf_bytes, content_type="application/pdf")
            if docx_bytes:
                key = f"doc-center/generated/{doc_id}.docx"
                docx_url = await upload_bytes(
                    key, docx_bytes,
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        except Exception:
            doc_id = str(uuid.uuid4())

        # Fetch employee info
        emp_name  = resolved.get("employee_name", "")
        emp_email = resolved.get("employee_email", "")
        if req.employee_id and not emp_name:
            emp = await db.employees.find_one({"_id": req.employee_id}) or {}
            emp_name  = emp.get("full_name", "")
            emp_email = emp.get("email", "")

        gen = DocGenerated(
            template_id=req.template_id,
            template_name=tmpl["name"],
            document_name=req.document_name,
            employee_id=req.employee_id,
            employee_name=emp_name,
            employee_email=emp_email,
            field_values=resolved,
            html_content=html_content,
            status=DocStatus.GENERATED,
            pdf_url=pdf_url,
            docx_url=docx_url,
            created_by=user_id,
            created_by_name=user_name,
        )
        gen_doc = gen.model_dump(by_alias=True)
        gen_doc["_id"] = doc_id
        await db.doc_generated.insert_one(gen_doc)

        # Increment generate_count on template
        await db.doc_templates.update_one(
            {"_id": req.template_id},
            {"$inc": {"generate_count": 1}, "$set": {"updated_at": datetime.now(timezone.utc)}},
        )

        # Attach bytes for inline download response
        gen_doc["_pdf_bytes"]  = pdf_bytes
        gen_doc["_docx_bytes"] = docx_bytes

        return True, "Document generated", gen_doc

    async def list_generated(
        self,
        db: AsyncIOMotorDatabase,
        *,
        template_id: Optional[str] = None,
        employee_id: Optional[str] = None,
        status: Optional[str] = None,
        search: Optional[str] = None,
        skip: int = 0,
        limit: int = 50,
    ) -> Tuple[List[Dict], int]:
        query: Dict[str, Any] = {"is_deleted": False}
        if template_id:  query["template_id"]  = template_id
        if employee_id:  query["employee_id"]   = employee_id
        if status:       query["status"]        = status
        if search:
            query["$or"] = [
                {"document_name":  {"$regex": search, "$options": "i"}},
                {"template_name":  {"$regex": search, "$options": "i"}},
                {"employee_name":  {"$regex": search, "$options": "i"}},
            ]
        total  = await db.doc_generated.count_documents(query)
        cursor = db.doc_generated.find(query, {"html_content": 0, "field_values": 0}).sort("created_at", -1).skip(skip).limit(limit)
        docs   = await cursor.to_list(length=limit)
        return docs, total

    async def delete_generated(self, db: AsyncIOMotorDatabase, doc_id: str) -> Tuple[bool, str]:
        result = await db.doc_generated.update_one(
            {"_id": doc_id},
            {"$set": {"is_deleted": True, "updated_at": datetime.now(timezone.utc)}},
        )
        return (True, "Deleted") if result.matched_count else (False, "Not found")

    # ── Approval Workflow ──────────────────────────────────────────────────────

    async def request_approval(
        self, db: AsyncIOMotorDatabase, data: DocApprovalCreate, user_id: str, user_name: str
    ) -> Tuple[bool, str, Optional[Dict]]:
        tmpl = await db.doc_templates.find_one({"_id": data.template_id})
        if not tmpl:
            return False, "Template not found", None
        # Move to review
        await db.doc_templates.update_one(
            {"_id": data.template_id},
            {"$set": {"status": DocStatus.REVIEW, "updated_at": datetime.now(timezone.utc)}},
        )
        appr = DocApproval(
            template_id=data.template_id,
            template_name=tmpl["name"],
            requested_by=user_id,
            requested_by_name=user_name,
            comments=data.comments,
        )
        doc = appr.model_dump(by_alias=True)
        await db.doc_approvals.insert_one(doc)
        return True, "Approval requested", doc

    async def list_approvals(
        self,
        db: AsyncIOMotorDatabase,
        *,
        status: Optional[str] = None,
        user_id: Optional[str] = None,
        skip: int = 0,
        limit: int = 50,
    ) -> Tuple[List[Dict], int]:
        query: Dict[str, Any] = {}
        if status:   query["status"]          = status
        if user_id:  query["requested_by"]    = user_id
        total  = await db.doc_approvals.count_documents(query)
        cursor = db.doc_approvals.find(query).sort("created_at", -1).skip(skip).limit(limit)
        docs   = await cursor.to_list(length=limit)
        return docs, total

    async def review_approval(
        self,
        db: AsyncIOMotorDatabase,
        approval_id: str,
        data: DocApprovalReview,
        user_id: str,
        user_name: str,
    ) -> Tuple[bool, str]:
        appr = await db.doc_approvals.find_one({"_id": approval_id})
        if not appr:
            return False, "Approval not found"
        if appr["status"] != ApprovalStatus.PENDING:
            return False, "Approval already reviewed"
        now = datetime.now(timezone.utc)
        await db.doc_approvals.update_one({"_id": approval_id}, {"$set": {
            "status": data.status,
            "approver_id": user_id,
            "approver_name": user_name,
            "reviewer_comments": data.reviewer_comments,
            "reviewed_at": now,
            "updated_at": now,
        }})
        new_status = DocStatus.APPROVED if data.status == ApprovalStatus.APPROVED else DocStatus.DRAFT
        await db.doc_templates.update_one(
            {"_id": appr["template_id"]},
            {"$set": {"status": new_status, "updated_at": now}},
        )
        return True, f"Template {data.status}"

    # ── Import ─────────────────────────────────────────────────────────────────

    async def import_file(
        self,
        db: AsyncIOMotorDatabase,
        filename: str,
        file_bytes: bytes,
        file_type: str,
        name: str,
        description: str,
        category_id: Optional[str],
        tags: List[str],
        user_id: str,
        user_name: str,
    ) -> Tuple[bool, str, Optional[Dict]]:
        # Extract text from uploaded file as body_html
        body_html = ""
        if file_type == "html":
            body_html = file_bytes.decode("utf-8", errors="replace")
        elif file_type == "docx":
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        style = para.style.name
                        if style.startswith("Heading"):
                            lvl = style.replace("Heading ", "")
                            parts.append(f"<h{lvl}>{para.text}</h{lvl}>")
                        else:
                            parts.append(f"<p>{para.text}</p>")
                body_html = "\n".join(parts)
            except Exception as e:
                logger.warning(f"DOCX import parse error: {e}")
                body_html = f"<p>{filename}</p>"
        elif file_type == "pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        text_parts.append(f"<p>{text.strip()}</p>")
                body_html = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PDF import parse error: {e}")
                body_html = f"<p>{filename}</p>"

        # Store raw file in S3
        s3_key = None
        try:
            from app.utils.s3 import upload_bytes
            s3_key = f"doc-center/imports/{uuid.uuid4()}/{filename}"
            await upload_bytes(s3_key, file_bytes, content_type=f"application/{file_type}")
        except Exception:
            pass

        content = TemplateContent(body_html=body_html)
        tmpl = DocTemplate(
            name=name,
            description=description,
            category_id=category_id,
            template_type=DocTemplateType.IMPORTED,
            tags=tags,
            content=content,
            s3_key=s3_key,
            original_filename=filename,
            file_type=file_type,
            created_by=user_id,
            created_by_name=user_name,
            updated_by=user_id,
        )
        doc = tmpl.model_dump(by_alias=True)
        await db.doc_templates.insert_one(doc)
        return True, "File imported successfully", doc

    # ── Template Library ───────────────────────────────────────────────────────

    def get_library_list(self) -> List[Dict]:
        return TEMPLATE_LIBRARY

    async def create_from_library(
        self,
        db: AsyncIOMotorDatabase,
        key: str,
        user_id: str,
        user_name: str,
        category_id: Optional[str] = None,
    ) -> Tuple[bool, str, Optional[Dict]]:
        tmpl_meta = next((t for t in TEMPLATE_LIBRARY if t["key"] == key), None)
        if not tmpl_meta:
            return False, "Template not found in library", None

        body_html = TEMPLATE_BODIES.get(key, f"<h2>{tmpl_meta['name']}</h2>\n<p>Add your content here...</p>")
        content = TemplateContent(body_html=body_html)

        data = DocTemplateCreate(
            name=tmpl_meta["name"],
            description=tmpl_meta["description"],
            category_id=category_id,
            template_type=DocTemplateType.SIMPLE,
            tags=[tmpl_meta["category"]],
            content=content,
            dynamic_fields=re.findall(r"\{\{(\w+)\}\}", body_html),
            change_summary="Created from library",
        )
        doc = await self.create_template(db, data, user_id, user_name)
        return True, "Template created from library", doc

    # ── Archive ────────────────────────────────────────────────────────────────

    async def list_archived(self, db: AsyncIOMotorDatabase, skip: int = 0, limit: int = 50) -> Tuple[List[Dict], int]:
        query = {"is_archived": True, "is_deleted": False}
        total  = await db.doc_templates.count_documents(query)
        cursor = db.doc_templates.find(query, {"content.body_html": 0}).sort("updated_at", -1).skip(skip).limit(limit)
        docs   = await cursor.to_list(length=limit)
        return docs, total

    async def unarchive_template(self, db: AsyncIOMotorDatabase, template_id: str) -> Tuple[bool, str]:
        result = await db.doc_templates.update_one(
            {"_id": template_id},
            {"$set": {"is_archived": False, "updated_at": datetime.now(timezone.utc)}},
        )
        return (True, "Unarchived") if result.matched_count else (False, "Not found")

    # ── Global Version History ─────────────────────────────────────────────────

    async def list_all_versions(
        self,
        db: AsyncIOMotorDatabase,
        *,
        template_id: Optional[str] = None,
        search: Optional[str] = None,
        skip: int = 0,
        limit: int = 50,
    ) -> Tuple[List[Dict], int]:
        query: Dict[str, Any] = {}
        if template_id:
            query["template_id"] = template_id
        if search:
            query["$or"] = [
                {"name":           {"$regex": search, "$options": "i"}},
                {"change_summary": {"$regex": search, "$options": "i"}},
                {"created_by_name":{"$regex": search, "$options": "i"}},
            ]
        total  = await db.doc_template_versions.count_documents(query)
        # Exclude heavy content blob from list view
        cursor = db.doc_template_versions.find(query, {"content": 0}).sort("created_at", -1).skip(skip).limit(limit)
        docs   = await cursor.to_list(length=limit)
        result = []
        for d in docs:
            d["id"] = str(d.pop("_id", ""))
            result.append(d)
        return result, total


document_center_service = DocumentCenterService()
